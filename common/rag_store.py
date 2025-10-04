# common/rag_store.py
import os, math, re, json, hashlib, time
from dataclasses import dataclass
from typing import List, Tuple, Optional
from .config import (
    DOCS_PATH, CHUNK_SIZE, CHUNK_OVERLAP, USE_OPENAI_EMBEDDINGS,
    OPENAI_EMBED_MODEL, EMBED_BATCH_SIZE, RAG_CACHE_DIR
)

# Optional OpenAI client (graceful fallback)
_client = None
if USE_OPENAI_EMBEDDINGS:
    try:
        from openai import OpenAI
        _client = OpenAI()
    except Exception:
        _client = None

from docx import Document
_WORD = re.compile(r"[A-Za-z0-9_]+")

def _read_docx(path: str) -> str:
    doc = Document(path)
    parts: List[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t: parts.append(t)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text): parts.append("\t".join(row_text))
    full = "\n".join(parts)
    full = re.sub(r"\n{3,}", "\n\n", full)
    return full

def _read_file(path: str) -> str:
    if not os.path.exists(path):
        return ""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        try:
            return _read_docx(path)
        except Exception:
            return ""
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""

def _chunk(text: str, size: int, overlap: int) -> List[str]:
    if not text: return []
    chunks, i = [], 0
    step = max(1, size - overlap)
    while i < len(text):
        chunks.append(text[i:i+size])
        i += step
    return chunks

# ---- sparse fallback ----
def _tokens(s: str) -> List[str]:
    return [t.lower() for t in _WORD.findall(s)]

def _bow(tokens: List[str]) -> dict:
    bag = {}
    for t in tokens:
        bag[t] = bag.get(t, 0) + 1
    return bag

def _normalize_sparse(vec: dict) -> dict:
    norm = math.sqrt(sum(v*v for v in vec.values())) or 1.0
    return {k: v / norm for k, v in vec.items()}

def _cos_sparse(a: dict, b: dict) -> float:
    if not a or not b: return 0.0
    if len(a) > len(b): a, b = b, a
    return sum(a[k]*b.get(k,0.0) for k in a)

# ---- dense ----
def _cos_dense(a: List[float], b: List[float]) -> float:
    if not a or not b: return 0.0
    s = sum(x*y for x, y in zip(a, b))
    na = math.sqrt(sum(x*x for x in a)) or 1.0
    nb = math.sqrt(sum(y*y for y in b)) or 1.0
    return s / (na * nb)

def _doc_signature(path: str) -> str:
    try:
        st = os.stat(path)
        payload = f"{path}|{st.st_size}|{int(st.st_mtime)}"
    except FileNotFoundError:
        payload = f"{path}|0|0"
    return hashlib.sha256(payload.encode()).hexdigest()

def _ensure_dir(p: str):
    os.makedirs(p, exist_ok=True)

@dataclass
class RagChunk:
    text: str
    meta: dict
    vec_sparse: Optional[dict] = None
    vec_dense: Optional[List[float]] = None

class RagStore:
    def __init__(self, path: str = DOCS_PATH):
        self.path = path
        self.chunks: List[RagChunk] = []
        self._build()

    def _build(self):
        text = _read_file(self.path)
        parts = _chunk(text, CHUNK_SIZE, CHUNK_OVERLAP)
        if USE_OPENAI_EMBEDDINGS and _client is not None:
            self._build_dense(parts)
        else:
            self._build_sparse(parts)

    def _build_sparse(self, parts: List[str]):
        self.chunks = []
        for i, p in enumerate(parts):
            vec = _normalize_sparse(_bow(_tokens(p)))
            self.chunks.append(RagChunk(text=p, meta={"chunk_id": i}, vec_sparse=vec))

    def _build_dense(self, parts: List[str]):
        _ensure_dir(RAG_CACHE_DIR)
        sig = _doc_signature(self.path)
        cache_file = os.path.join(RAG_CACHE_DIR, f"{sig}.embeddings.json")
        if os.path.exists(cache_file):
            try:
                with open(cache_file, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if data.get("model") == OPENAI_EMBED_MODEL and len(data.get("chunks", [])) == len(parts):
                    self.chunks = [
                        RagChunk(text=entry["text"], meta={"chunk_id": i}, vec_dense=entry["vec"])
                        for i, entry in enumerate(data["chunks"])
                    ]
                    return
            except Exception:
                pass
        vectors: List[List[float]] = []
        for i in range(0, len(parts), EMBED_BATCH_SIZE):
            batch = parts[i:i+EMBED_BATCH_SIZE]
            resp = _client.embeddings.create(model=OPENAI_EMBED_MODEL, input=batch)
            vectors.extend([d.embedding for d in resp.data])
        self.chunks = [
            RagChunk(text=p, meta={"chunk_id": i}, vec_dense=vectors[i])
            for i, p in enumerate(parts)
        ]
        payload = {"model": OPENAI_EMBED_MODEL, "created": int(time.time()),
                   "chunks": [{"text": c.text, "vec": c.vec_dense} for c in self.chunks]}
        with open(cache_file, "w", encoding="utf-8") as f:
            json.dump(payload, f)

    def retrieve(self, query: str, k: int = 5) -> List[Tuple[RagChunk, float]]:
        if USE_OPENAI_EMBEDDINGS and _client is not None and self.chunks and self.chunks[0].vec_dense is not None:
            q = _client.embeddings.create(model=OPENAI_EMBED_MODEL, input=[query]).data[0].embedding
            scored = [(c, _cos_dense(q, c.vec_dense)) for c in self.chunks]
        else:
            q = _normalize_sparse(_bow(_tokens(query)))
            scored = [(c, _cos_sparse(q, c.vec_sparse or {})) for c in self.chunks]
        scored.sort(key=lambda x: x[1], reverse=True)
        return scored[:k]

_store = None
def get_store():
    global _store
    if _store is None:
        _store = RagStore(DOCS_PATH)
    return _store
